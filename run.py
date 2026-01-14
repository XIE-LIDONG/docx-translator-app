import streamlit as st
import tempfile
from docx import Document
from deep_translator import GoogleTranslator
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
import os
import re
import io

# Page configuration
st.set_page_config(page_title="DOCX Document Language Translator", page_icon="📄", layout="wide")

# 标题+署名：同一行布局（标题左，署名右）
st.markdown(
    """
    <div style='display: flex; justify-content: space-between; align-items: center;'>
        <h1 style='margin: 0;'>📄 DOCX Document Language Translator</h1>
        <p style='color: #666666; font-size: 14px; margin: 0;'>By XIE LI DONG</p>
    </div>
    """,
    unsafe_allow_html=True
)

# Added usage tip below title
st.info("""
💡 **Usage Tip**: If your file is in PDF format, convert it to DOCX first via [ILovePDF](https://www.ilovepdf.com/). 
After translation, you can convert the DOCX back to PDF using ILovePDF if needed. 
We tried integrating PDF-DOCX conversion directly into Streamlit, but it drastically slowed down the entire application.
""")
st.markdown("---")

# Define supported languages (Name: deep-translator code)
SUPPORT_LANGUAGES = {
    "Chinese": "zh-CN",
    "English": "en",
    "French": "fr",
    "German": "de",
    "Spanish": "es",
    "Arabic": "ar",
    "Japanese": "ja",
    "Korean": "ko"
}
# Extract language name list (for dropdown)
LANG_NAMES = list(SUPPORT_LANGUAGES.keys())

# File upload
uf = st.file_uploader("Select Word Document (.docx)", type=["docx"])

if uf:
    # File information
    c1, c2, c3 = st.columns([2, 1, 1])
    with c1: 
        st.success(f"📁 **File:** {uf.name}")
    with c2: 
        upload_size = uf.size/(1024*1024)
        st.metric("Upload Size", f"{upload_size:.2f} MB")
    with c3:
        # 预先读取文件内容
        uf_content = uf.getvalue()
        st.metric("Content Size", f"{len(uf_content)/(1024*1024):.2f} MB")
    st.markdown("---")

    # Translation settings (multilingual dropdown + 线程/批次配置)
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        source_lang_name = st.selectbox(
            "**Source Language**",
            LANG_NAMES,
            index=LANG_NAMES.index("English")  # 根据您的文档，源语言是英文
        )
        source_lang = SUPPORT_LANGUAGES[source_lang_name]
    with c2:
        target_lang_name = st.selectbox(
            "**Target Language**",
            LANG_NAMES,
            index=LANG_NAMES.index("Chinese")  # 翻译为中文
        )
        target_lang = SUPPORT_LANGUAGES[target_lang_name]
    with c3:
        wk = st.slider(
            "**Thread Count**",
            min_value=1,
            max_value=3,
            value=2,
            help="Number of parallel translation threads (1-3 for stability)"
        )
    with c4:
        BS = st.slider(
            "**Batch Size**",
            min_value=20,
            max_value=100,
            value=50,
            step=10,
            help="Number of text segments per translation batch (20-100)"
        )

    if st.button("🚀 Start Translation", type="primary", use_container_width=True):
        # Progress log area
        log_area = st.empty()
        log = []
        
        # 用于存储所有要处理的元素
        text_items = []  # (原始对象, 原始文本, 类型, 额外信息)
        all_texts = []   # 所有文本片段
        doc = None
        fp = None

        try:
            stt = time.time()
            
            # ========== 1. 保存临时文件 ==========
            with io.BytesIO(uf_content) as bio:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", mode='wb+') as tmp:
                    tmp.write(bio.getvalue())
                    tmp.flush()
                    tmp.seek(0)
                    fp = tmp.name
            
            # ========== 2. 解析文档 ==========
            doc = Document(fp)
            log.append(f"✅ DOCX parsed: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
            log_area.markdown("\n".join(log))
            
            # ========== 3. 增强的文本提取：使用正则表达式提取所有文本 ==========
            log.append("🔍 Extracting all text content...")
            log_area.markdown("\n".join(log))
            
            # 3.1 提取段落文本
            paragraph_count = 0
            for para_idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    # 使用正则表达式处理特殊格式
                    # 去除多余的空格和换行符
                    cleaned_text = re.sub(r'\s+', ' ', text).strip()
                    
                    # 检查是否是表格样式的文本（包含分隔线）
                    if re.search(r'[-+]{3,}', cleaned_text) or (re.search(r'\|\s*[-\w]', cleaned_text) and len(cleaned_text.split('\n')) > 1):
                        # 按行分割并处理每一行
                        lines = cleaned_text.split('\n')
                        for line in lines:
                            line = line.strip()
                            # 去除表格分隔符
                            line = re.sub(r'^[-+|]+$', '', line)
                            line = re.sub(r'^\s*[-+|\s]+\s*$', '', line)
                            if line:
                                text_items.append((para, line, 'paragraph_line', {'para_idx': para_idx, 'line_idx': len(text_items)}))
                                all_texts.append(line)
                                paragraph_count += 1
                    else:
                        # 普通段落文本
                        text_items.append((para, cleaned_text, 'paragraph', {'para_idx': para_idx}))
                        all_texts.append(cleaned_text)
                        paragraph_count += 1
                    
                    # 每提取50个段落显示一次进度
                    if paragraph_count % 50 == 0:
                        log.append(f"📝 Extracted {paragraph_count} paragraph segments")
                        log_area.markdown("\n".join(log))
            
            # 3.2 提取表格文本（增强版）
            table_count = 0
            cell_count = 0
            
            for table_idx, table in enumerate(doc.tables):
                log.append(f"📊 Processing table {table_idx + 1}/{len(doc.tables)}")
                log_area.markdown("\n".join(log))
                
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        # 提取单元格的所有文本内容
                        cell_text = cell.text.strip()
                        
                        if cell_text:
                            # 使用正则表达式清理单元格文本
                            # 去除多余的空格和换行符
                            cleaned_cell_text = re.sub(r'\s+', ' ', cell_text).strip()
                            
                            # 如果单元格内有多个段落，分割处理
                            if '\n' in cleaned_cell_text or '\r' in cleaned_cell_text:
                                # 按换行符分割
                                lines = re.split(r'[\n\r]+', cleaned_cell_text)
                                for line_idx, line in enumerate(lines):
                                    line = line.strip()
                                    if line:
                                        text_items.append(
                                            (cell, line, 'cell_line', 
                                             {'table_idx': table_idx, 'row_idx': row_idx, 
                                              'cell_idx': cell_idx, 'line_idx': line_idx})
                                        )
                                        all_texts.append(line)
                                        cell_count += 1
                            else:
                                # 单行单元格文本
                                text_items.append(
                                    (cell, cleaned_cell_text, 'cell', 
                                     {'table_idx': table_idx, 'row_idx': row_idx, 'cell_idx': cell_idx})
                                )
                                all_texts.append(cleaned_cell_text)
                                cell_count += 1
                        
                        # 每提取30个单元格显示一次进度
                        if cell_count % 30 == 0:
                            log.append(f"📋 Extracted {cell_count} table cell segments")
                            log_area.markdown("\n".join(log))
                
                table_count += 1
            
            # 3.3 提取页眉页脚文本
            try:
                # 尝试提取页眉
                for section in doc.sections:
                    header = section.header
                    for para in header.paragraphs:
                        text = para.text.strip()
                        if text:
                            cleaned_text = re.sub(r'\s+', ' ', text).strip()
                            text_items.append((para, cleaned_text, 'header', {}))
                            all_texts.append(cleaned_text)
                
                log.append(f"📄 Extracted header/footer text")
            except:
                log.append("⚠️ Header/footer extraction not supported or no headers/founders")
            
            # 3.4 汇总提取结果
            total_segments = len(all_texts)
            log.append(f"✅ Extraction completed: {total_segments} total segments")
            log.append(f"   - Paragraph segments: {paragraph_count}")
            log.append(f"   - Table cell segments: {cell_count}")
            log.append(f"   - Header segments: {total_segments - paragraph_count - cell_count}")
            
            # 显示一些示例文本（调试用）
            if total_segments > 0:
                sample_count = min(5, total_segments)
                log.append(f"📋 Sample of extracted texts (first {sample_count}):")
                for i in range(sample_count):
                    text_preview = all_texts[i][:80] + ('...' if len(all_texts[i]) > 80 else '')
                    log.append(f"   {i+1}. [{text_items[i][2]}] {text_preview}")
            
            log_area.markdown("\n".join(log))
            
            # 检查是否有文本可翻译
            if total_segments == 0:
                st.error("❌ No valid text found in document")
                st.stop()
            
            # ========== 4. 多线程翻译 ==========
            log.append(f"🔤 Translation direction: {source_lang_name} → {target_lang_name}")
            log.append(f"⚙️ Configuration: {wk} threads | {BS} segments per batch")
            log_area.markdown("\n".join(log))
            
            translations = [None] * total_segments
            
            def translate_batch(batch_texts):
                """翻译批处理函数"""
                try:
                    # 卡车术语预处理（针对您的文档优化）
                    term_map = {
                        "GVW": "Gross Vehicle Weight",
                        "Curb Weight": "Curb Weight of Chassis",
                        "Axle Load": "Axle Load Distribution",
                        "Wheel Base": "Wheel Base Length",
                        "Max Torque": "Maximum Torque",
                        "Fuel Consumption": "Fuel Consumption per 100km",
                        "ABS": "Anti-lock Braking System",
                        "B10": "B10 life",
                        "GCC": "Gulf Cooperation Council",
                        "KSA": "Kingdom of Saudi Arabia",
                        "WABCO": "WABCO brand",
                        "EATON": "EATON brand"
                    }
                    
                    # 处理术语
                    processed_texts = []
                    for text in batch_texts:
                        processed = text
                        for term, full in term_map.items():
                            # 使用正则表达式确保完整单词匹配
                            pattern = r'\b' + re.escape(term) + r'\b'
                            processed = re.sub(pattern, full, processed)
                        processed_texts.append(processed)
                    
                    # 翻译
                    translated = GoogleTranslator(
                        source=source_lang, 
                        target=target_lang
                    ).translate_batch(processed_texts)
                    
                    # 检查翻译结果
                    result = []
                    for i, trans in enumerate(translated):
                        if trans and trans.strip():
                            # 还原术语
                            final = trans
                            for term, full in term_map.items():
                                pattern = r'\b' + re.escape(full) + r'\b'
                                final = re.sub(pattern, term, final)
                            result.append(final)
                        else:
                            # 翻译失败时使用原文
                            result.append(batch_texts[i])
                    
                    return result
                except Exception as e:
                    log.append(f"⚠️ Batch translation error: {str(e)[:50]}")
                    # 出错时返回原文
                    return batch_texts
            
            # 多线程翻译
            with ThreadPoolExecutor(max_workers=wk) as executor:
                futures = {}
                
                # 提交批次任务
                for start_idx in range(0, total_segments, BS):
                    end_idx = min(start_idx + BS, total_segments)
                    batch = all_texts[start_idx:end_idx]
                    future = executor.submit(translate_batch, batch)
                    futures[future] = (start_idx, end_idx)
                
                # 处理完成的任务
                completed = 0
                for future in as_completed(futures):
                    start_idx, end_idx = futures[future]
                    try:
                        batch_result = future.result()
                        # 保存翻译结果
                        for i, trans in enumerate(batch_result):
                            if start_idx + i < total_segments:
                                translations[start_idx + i] = trans
                        
                        completed += len(batch_result)
                        # 显示进度
                        if completed % 20 == 0 or completed == total_segments:
                            log.append(f"🔄 Translated: {completed}/{total_segments} segments")
                            log_area.markdown("\n".join(log))
                    except Exception as e:
                        log.append(f"❌ Error processing batch: {str(e)[:50]}")
                        log_area.markdown("\n".join(log))
            
            log.append(f"✅ Translation completed: {total_segments}/{total_segments}")
            log_area.markdown("\n".join(log))
            
            # ========== 5. 更新文档 ==========
            log.append("📝 Updating document with translations...")
            log_area.markdown("\n".join(log))
            
            update_count = 0
            
            # 根据类型更新文档
            for idx, (obj, original_text, obj_type, extra_info) in enumerate(text_items):
                if idx < len(translations) and translations[idx]:
                    translated_text = translations[idx]
                    
                    if obj_type in ['paragraph', 'paragraph_line', 'header']:
                        # 更新段落
                        try:
                            obj.text = translated_text
                            update_count += 1
                        except:
                            # 如果无法直接设置text属性，尝试其他方式
                            try:
                                for run in obj.runs:
                                    run.text = ""
                                if obj.runs:
                                    obj.runs[0].text = translated_text
                                else:
                                    obj.add_run(translated_text)
                                update_count += 1
                            except:
                                log.append(f"⚠️ Could not update paragraph at index {idx}")
                    
                    elif obj_type in ['cell', 'cell_line']:
                        # 更新表格单元格
                        try:
                            obj.text = translated_text
                            update_count += 1
                        except:
                            # 如果无法直接设置text属性
                            try:
                                # 清空单元格内容
                                for paragraph in obj.paragraphs:
                                    for run in paragraph.runs:
                                        run.text = ""
                                # 添加新内容
                                if obj.paragraphs:
                                    obj.paragraphs[0].text = translated_text
                                update_count += 1
                            except:
                                log.append(f"⚠️ Could not update table cell at index {idx}")
            
            log.append(f"✅ Document updated: {update_count}/{total_segments} segments modified")
            log_area.markdown("\n".join(log))
            
            # ========== 6. 保存并下载 ==========
            output_filename = f"{os.path.splitext(uf.name)[0]}_{source_lang_name}2{target_lang_name}.docx"
            output_path = fp.replace(".docx", f"_{source_lang_name}2{target_lang_name}.docx")
            
            doc.save(output_path)
            total_time = time.time() - stt
            
            # 成功提示
            st.balloons()
            st.success(f"### ✅ Translation Completed! ({source_lang_name} → {target_lang_name})")
            
            # 显示统计信息
            col1, col2, col3 = st.columns(3)
            col1.metric("Total Time", f"{total_time:.1f}s")
            col2.metric("Total Segments", f"{total_segments}")
            col3.metric("Updated Segments", f"{update_count}")
            
            # 显示处理速度
            if total_time > 0:
                speed = total_segments / total_time
                st.info(f"📊 Processing speed: {speed:.1f} segments/second")
            
            # 下载按钮
            st.markdown("---")
            with open(output_path, "rb") as f:
                st.download_button(
                    "📥 Download Translated Document",
                    f,
                    file_name=output_filename,
                    use_container_width=True,
                    type="primary"
                )
            
        except Exception as e:
            st.error("### ❌ Translation Failed")
            st.exception(e)
            if 'log' in locals():
                log.append(f"❌ Error: {str(e)[:100]}")
                log_area.markdown("\n".join(log))
        
        finally:
            # 清理临时文件
            try:
                if fp and os.path.exists(fp):
                    os.unlink(fp)
                if 'output_path' in locals() and output_path and os.path.exists(output_path):
                    os.unlink(output_path)
                log.append("✅ Temporary files cleaned up")
                if 'log_area' in locals():
                    log_area.markdown("\n".join(log))
            except Exception as cleanup_e:
                st.warning(f"⚠️ Cleanup warning: {str(cleanup_e)[:50]}...")
